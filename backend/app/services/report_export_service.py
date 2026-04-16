"""
正式报告导出：PDF（reportlab + 内嵌图表）、Word（python-docx）。

本次修复：
- matplotlib 在容器内常因缺字体/无 DISPLAY 而导致 savefig 卡死或渲染失败：
  * 强制 backend=Agg（import 之前就 `matplotlib.use("Agg")`）
  * 关掉交互模式，显式关闭字体警告
  * 尝试按顺序加载多个常见 CJK 字体；都不可用时降级为 DejaVu（仅 ASCII 标签）
- 所有 `plt.close` / `fig.savefig` 均做异常兜底，保证主流程（PDF/Word）不因图表失败而中断。
"""

from __future__ import annotations

import io
import logging
from typing import Any

logger = logging.getLogger(__name__)


_MPL_READY = False


def _ensure_matplotlib_ready() -> None:
    """第一次使用前：设置无头 Agg backend + CJK 字体，避免 worker 内卡死。"""
    global _MPL_READY
    if _MPL_READY:
        return
    try:
        import matplotlib

        matplotlib.use("Agg", force=True)
        import matplotlib.pyplot as plt
        from matplotlib import font_manager, rcParams

        rcParams["axes.unicode_minus"] = False
        rcParams["figure.max_open_warning"] = 0
        plt.ioff()

        # 按优先级探测系统中可用的 CJK 字体；任何一种可用即可
        candidates = [
            "Noto Sans CJK SC",
            "Noto Sans CJK",
            "Source Han Sans SC",
            "Source Han Sans CN",
            "WenQuanYi Zen Hei",
            "WenQuanYi Micro Hei",
            "PingFang SC",
            "Microsoft YaHei",
            "SimHei",
            "Arial Unicode MS",
        ]
        available = {f.name for f in font_manager.fontManager.ttflist}
        picked = next((c for c in candidates if c in available), None)
        if picked:
            rcParams["font.sans-serif"] = [picked] + list(rcParams.get("font.sans-serif", []))
        else:
            logger.info("report_chart: 未检测到 CJK 字体，图表标签将使用 ASCII 文案兜底")
        _MPL_READY = True
    except Exception:
        logger.exception("matplotlib_init_failed — chart rendering will be skipped")
        _MPL_READY = False


def _chart_png_bytes(payload: dict[str, Any]) -> bytes:
    """简单柱状图（风险维度），避免中文轴标签字体问题。"""
    _ensure_matplotlib_ready()
    import matplotlib.pyplot as plt

    bi = payload.get("basic_info") or {}
    eco = payload.get("economic") or {}
    try:
        score = float(bi.get("risk_score") or 0)
    except (TypeError, ValueError):
        score = 0.0
    try:
        amt = float(eco.get("total_amount") or 0)
    except (TypeError, ValueError):
        amt = 0.0
    try:
        ar = float(eco.get("anomaly_ratio") or 0) * 100
    except (TypeError, ValueError):
        ar = 0.0

    fig, ax = plt.subplots(figsize=(6, 3.2), dpi=120)
    try:
        labels = ["risk_score", "anomaly_%", "amount_k"]
        vals = [
            max(0.0, min(100.0, score)),
            max(0.0, min(100.0, ar)),
            max(0.0, min(100.0, (amt / 1_000_000) * 10)) if amt > 0 else 0.0,
        ]
        ax.bar(labels, vals, color=["#4472c4", "#ed7d31", "#a5a5a5"])
        ax.set_ylim(0, 105)
        ax.set_title("Key metrics (normalized)")
        for i, v in enumerate(vals):
            ax.text(i, v + 2, f"{v:.1f}", ha="center", va="bottom", fontsize=9)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png")
        buf.seek(0)
        return buf.read()
    finally:
        try:
            plt.close(fig)
        except Exception:
            logger.debug("plt_close_failed", exc_info=True)


def render_pdf_bytes(payload: dict[str, Any]) -> bytes:
    """PDF：基本信息、风险、线索表、图表。"""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    try:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    except Exception:
        logger.debug("register_cid_font_failed", exc_info=True)

    bio = io.BytesIO()
    doc = SimpleDocTemplate(
        bio, pagesize=A4,
        rightMargin=2 * cm, leftMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "t",
        parent=styles["Heading1"],
        fontName="STSong-Light",
        fontSize=18,
        spaceAfter=12,
    )
    body = ParagraphStyle(
        "b",
        parent=styles["Normal"],
        fontName="STSong-Light",
        fontSize=11,
        leading=16,
    )
    story: list[Any] = []

    story.append(Paragraph("侦查辅助分析报告", title_style))
    story.append(Spacer(1, 0.4 * cm))

    bi = payload.get("basic_info") or {}
    story.append(Paragraph(f"<b>案件编号</b>：{bi.get('case_id', '')}", body))
    story.append(Paragraph(f"<b>对象标识</b>：{bi.get('person_id', '')}", body))
    story.append(Paragraph(
        f"<b>综合风险</b>：{bi.get('risk_score', 0)}（{bi.get('risk_level', '')}）", body,
    ))
    story.append(Paragraph(f"<b>摘要</b>：{bi.get('summary', '')}", body))
    story.append(Spacer(1, 0.5 * cm))

    story.append(Paragraph("<b>一、经济状况</b>", body))
    eco = payload.get("economic") or {}
    story.append(
        Paragraph(
            f"估算总交易额：{eco.get('total_amount', 0):,.0f} 元；"
            f"异常线索占比：{float(eco.get('anomaly_ratio') or 0) * 100:.1f}%；"
            f"转出 {eco.get('transfer_out_count', 0)} 笔，转入 {eco.get('transfer_in_count', 0)} 笔。"
            f"{eco.get('explain', '')}",
            body,
        )
    )
    story.append(Spacer(1, 0.3 * cm))

    story.append(Paragraph("<b>二、关键线索</b>", body))
    clues = payload.get("clues") or []
    if clues:
        data = [["标题", "类别", "风险等级", "评分"]]
        for c in clues[:20]:
            data.append([
                str(c.get("title", ""))[:40],
                str(c.get("category", "")),
                str(c.get("risk_level", "")),
                f"{float(c.get('risk_score') or 0):.0f}",
            ])
        t = Table(data, colWidths=[6 * cm, 2.5 * cm, 2.5 * cm, 2 * cm])
        t.setStyle(
            TableStyle([
                ("FONTNAME", (0, 0), (-1, -1), "STSong-Light"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#4472c4")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ])
        )
        story.append(t)
    else:
        story.append(Paragraph("暂无线索记录。", body))

    story.append(Spacer(1, 0.5 * cm))
    story.append(Paragraph("<b>三、指标图示</b>", body))
    try:
        img_data = _chart_png_bytes(payload)
        img = Image(io.BytesIO(img_data), width=14 * cm, height=7 * cm)
        story.append(img)
    except Exception:
        logger.exception("chart_embed_failed")
        story.append(Paragraph("（图表生成失败，略）", body))

    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("<i>本报告由系统自动生成，仅供办案参考。</i>", body))

    doc.build(story)
    return bio.getvalue()


def render_docx_bytes(payload: dict[str, Any]) -> bytes:
    """Word：与 PDF 信息一致。"""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    d = docx.Document()
    t = d.add_heading("侦查辅助分析报告", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    bi = payload.get("basic_info") or {}
    p = d.add_paragraph()
    p.add_run("案件编号：").bold = True
    p.add_run(str(bi.get("case_id", "")))
    p = d.add_paragraph()
    p.add_run("对象标识：").bold = True
    p.add_run(str(bi.get("person_id", "")))
    p = d.add_paragraph()
    p.add_run("综合风险：").bold = True
    p.add_run(f"{bi.get('risk_score', 0)}（{bi.get('risk_level', '')}）")
    d.add_paragraph(str(bi.get("summary", "")))

    d.add_heading("一、经济状况", level=1)
    eco = payload.get("economic") or {}
    d.add_paragraph(
        f"估算总交易额 {eco.get('total_amount', 0):,.0f} 元；"
        f"异常线索占比 {float(eco.get('anomaly_ratio') or 0) * 100:.1f}%；"
        f"转出 {eco.get('transfer_out_count', 0)} 笔，转入 {eco.get('transfer_in_count', 0)} 笔。"
        f"{eco.get('explain', '')}"
    )

    d.add_heading("二、关键线索", level=1)
    clues = payload.get("clues") or []
    table = d.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "标题"
    hdr[1].text = "类别"
    hdr[2].text = "风险等级"
    hdr[3].text = "评分"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(10)
    for c in clues[:30]:
        row = table.add_row().cells
        row[0].text = str(c.get("title", ""))[:200]
        row[1].text = str(c.get("category", ""))
        row[2].text = str(c.get("risk_level", ""))
        row[3].text = f"{float(c.get('risk_score') or 0):.0f}"

    d.add_heading("三、行为与关系（摘要）", level=1)
    beh = payload.get("behavior") or {}
    d.add_paragraph(beh.get("explain", "") or "（详见系统内轨迹与关系子图）")
    soc = payload.get("social") or {}
    d.add_paragraph(soc.get("explain", "") or "")

    try:
        img_data = _chart_png_bytes(payload)
        d.add_paragraph()
        d.add_picture(io.BytesIO(img_data), width=Cm(14))
    except Exception:
        logger.exception("docx_chart_failed")

    foot = d.add_paragraph()
    fr = foot.add_run("本报告由系统自动生成，仅供办案参考。")
    fr.italic = True

    out = io.BytesIO()
    d.save(out)
    return out.getvalue()
